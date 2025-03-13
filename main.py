import streamlit as st
from docx import Document
from openai import OpenAI
from io import BytesIO
import os

# Streamlit User Interface
st.title("📌 Code Evaluator")
st.sidebar.write("🔑 **Get an OpenAI API Key:**")
st.sidebar.markdown("[Click here to get an API key](https://platform.openai.com/chat-completions)", unsafe_allow_html=True)

# Sidebar for API key input
api_key = st.sidebar.text_input("Enter your API key", type="password")

if api_key:
    st.sidebar.write("API Key Entered ✅")

# Default code (for testing)
code = """
Question 1 
Write a Python program to add two numbers?
code:
def add_numbers(a, b):
    return a + b

# Example usage
num1 = 5
num2 = 7
result = add_numbers(num1, num2)
print("Sum:", result)
"""

# Function to extract code from the DOCX file
def extract_code_from_docx(uploaded_file):
    """Extracts code from a DOCX file uploaded in Streamlit."""
    doc = Document(uploaded_file)
    code_blocks = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n".join(code_blocks)

# Upload file
uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

if uploaded_file:
    code = extract_code_from_docx(uploaded_file)

    if api_key:
        try:
            # ✅ Corrected OpenAI Client Initialization
            client = OpenAI(api_key=api_key)

            # Send code to OpenAI for evaluation
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                temperature=0.0,  # Ensures consistent responses
                messages=[
                    {"role": "system", "content": "You are an expert code evaluator."},
                    {"role": "user", "content": f"""
                    Task: Code Evaluation  

                     1. **Error Detection:**  
                     - Identify any syntax errors, logical errors, or runtime errors in the code.  
                     2. **Corrections & Improvements:**  
                     - Suggest corrections for the identified errors.
                     - Provide optimized alternatives where necessary.  
                      4. **Final Marking:**  
                      - Provide a final score out of 10.  
                      **Rules for Evaluation:**  
                      - Do not modify the logic of the code.  
                      - Keep feedback structured and consistent across evaluations.  
                      - Ensure suggestions are clear, repeatable, and deterministic.  
                      Evaluate the following code and provide detailed feedback:
                    ```python
                    {code}
                    ```
                    """}
                ]
            )

            # Show the result
            evaluation = response.choices[0].message.content
            st.subheader("📋 Code Evaluation:")
            st.write(evaluation)

            # Function to create a Word file from text
            def save_text_to_docx(text, uploaded_filename):
                """Converts evaluation text into a downloadable DOCX file."""
                doc = Document()
                doc.add_paragraph(text)
                byte_io = BytesIO()
                doc.save(byte_io)
                byte_io.seek(0)
                return byte_io, f"{uploaded_filename}_evaluation.docx"

            # Create Word file for download
            uploaded_name = os.path.splitext(uploaded_file.name)[0]  # Extract filename without extension
            word_file, filename = save_text_to_docx(evaluation, uploaded_name)

            st.download_button(
                label="📥 Download Evaluation as Word File",
                data=word_file,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"⚠️ Error: {e}")

            # Additional debugging messages
            if "authentication" in str(e).lower():
                st.warning("🔑 Check your OpenAI API key—it might be incorrect or expired.")
            elif "file" in str(e).lower():
                st.warning("📂 There might be an issue with the uploaded DOCX file.")
            else:
                st.warning("⚡ Try again or contact support.")
    else:
        st.error("❌ Please enter a valid API key to proceed.")
